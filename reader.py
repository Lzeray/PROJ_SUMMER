import csv
from pathlib import Path
import PyPDF2
import docx
import pandas as pd
from charset_normalizer import from_path
import docx2txt
import subprocess
import textract

class Reader:
    file_extensions = [
        '.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.csv'
    ]

    @classmethod
    def read_pdf(cls, file_path: Path) -> str:
        """Извлекает текст из PDF файла"""
        text = []
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
        except Exception as e:
            print(f"Ошибка при чтении PDF ({file_path.name}): {e}")
        return "\n".join(text) + "\n"
    
    @classmethod
    def read_docx(cls, file_path: Path) -> str:
        """Извлекает текст из DOCX файла"""
        text = []
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text.append(" | ".join(row_text))
        except Exception as e:
            print(f"Ошибка при чтении DOCX ({file_path.name}): {e}")
        return "\n".join(text) + "\n"
    
    @classmethod
    def read_doc(cls, file_path: Path) -> str:
        """Извлекает текст из DOC файла"""
        try:
            return docx2txt.process(file_path) + "\n"
        except:
            pass
        
        try:
            return textract.process(file_path).decode('utf-8') + "\n"
        except:
            pass
        
        try:
            result = subprocess.run(['catdoc', str(file_path)], capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout:
                return result.stdout + "\n"
        except:
            pass
            
    @classmethod
    def read_excel(cls, file_path: Path) -> str:
        """Извлекает данные из Excel файла (.xlsx, .xls)"""
        text = []
        try:
            engine = 'xlrd' if file_path.suffix.lower() == '.xls' else 'openpyxl'
            
            excel_file = pd.ExcelFile(file_path, engine=engine)
            for sheet_name in excel_file.sheet_names:
                text.append(f"\n--- Лист: {sheet_name} ---")
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine=engine)
                df_clean = df.dropna(how='all').fillna('')
                text.append(df_clean.to_string(index=False))
        except Exception as e:
            print(f"Ошибка при чтении Excel ({file_path.name}): {e}")
        return "\n".join(text) + "\n"
    
    @classmethod
    def read_text_file(cls, file_path: Path) -> str:
        """Автоматически определяет кодировку и читает текстовый файл"""
        try:
            result = from_path(file_path).best()
            if result:
                return str(result)
        except Exception as e:
            print(f"Ошибка при чтении текстового файла ({file_path.name}): {e}")
        return ""

    @classmethod
    def read_csv(cls, file_path: Path) -> str:
        """Автоматически определяет кодировку и читает CSV файл"""
        text = []
        try:
            blob = from_path(file_path).best()
            encoding = blob.encoding if blob else 'utf-8'
            
            with open(file_path, 'r', encoding=encoding) as file:
                sample = file.read(2048)
                file.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample)
                    reader = csv.reader(file, dialect)
                except csv.Error:
                    reader = csv.reader(file)
                    
                for row in reader:
                    if row:
                        text.append(", ".join(row))
        except Exception as e:
            print(f"Ошибка при чтении CSV ({file_path.name}): {e}")
        return "\n".join(text) + "\n"
    
    @classmethod
    def read_document(cls, file_path: str) -> str:
        """Чтение любого документа и вывод всего текста."""
        path = Path(file_path.strip('"\''))

        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        ext = path.suffix.lower()

        if ext == '.pdf':
            return cls.read_pdf(path)
        elif ext == '.docx':
            return cls.read_docx(path)
        elif ext == '.doc':
            return cls.read_doc(path)
        elif ext in ['.xlsx', '.xls']:
            return cls.read_excel(path)
        elif ext == '.txt':
            return cls.read_text_file(path)
        elif ext == '.csv':
            return cls.read_csv(path)
        else:
            return f"Неподдерживаемый формат файла: {ext}"
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Проверяет, поддерживается ли файл"""
        path = Path(file_path.strip('"\''))
        return path.suffix.lower() in cls.file_extensions